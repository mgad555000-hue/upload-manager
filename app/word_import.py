"""
Word Batch Import — تحليل ملفات Word واستيراد المواضيع
تنسيق الملف: جدول واحد في .docx
أول صف = عناوين الأعمدة
أعمدة أساسية: رقم، عنوان
أعمدة المنصات: "اسم المنصة - اسم الحقل" (مثال: يوتيوب - العنوان)
"""
import json
from typing import List, Dict, Tuple


def parse_word_file(
    file_path: str,
    channel_id: int,
    content_type: str,
    platform_map: Dict[str, int],
    field_map: Dict[int, Dict[str, str]],
) -> Tuple[List[Dict], List[str]]:
    """
    Parse a Word document table into topic data.

    Args:
        file_path: path to .docx file
        channel_id: target channel ID
        content_type: "shorts" or "long"
        platform_map: {display_name_or_name: platform_id}
        field_map: {platform_id: {field_label_or_name: field_name}}

    Returns:
        (topics_list, errors_list)
    """
    from docx import Document

    doc = Document(file_path)

    if not doc.tables:
        # Auto-detect MG Ranner format and use its parser
        try:
            from app.mg_ranner_import import parse_mg_ranner_docx
            mg_topics, mg_errors = parse_mg_ranner_docx(file_path)
            if mg_topics:
                # Convert MG Ranner parsed topics to word import format
                topics = []
                for i, mt in enumerate(mg_topics, start=1):
                    platform_data_list = []
                    fields = mt.get("fields", {})
                    # Group fields by platform
                    plat_fields_map = {}  # {platform_id: {field_name: value}}
                    for fkey, fval in fields.items():
                        if not fval:
                            continue
                        # Map MG Ranner field keys to platform fields
                        mapped = _map_mg_field(fkey, fval, platform_map, field_map)
                        if mapped:
                            pid, fname, fvalue = mapped
                            if pid not in plat_fields_map:
                                plat_fields_map[pid] = {}
                            plat_fields_map[pid][fname] = fvalue
                    for pid, fvals in plat_fields_map.items():
                        platform_data_list.append({
                            "platform_id": pid,
                            "field_values": fvals,
                        })
                    # Use yt_title_1 as topic title if no explicit title
                    title = mt.get("title") or fields.get("yt_title_1") or fields.get("yt_title")
                    topics.append({
                        "channel_id": channel_id,
                        "topic_number": mt.get("topic_number", i),
                        "content_type": content_type,
                        "title": title,
                        "platform_data": platform_data_list if platform_data_list else None,
                    })
                return topics, mg_errors
        except Exception:
            pass
        return [], ["\u0627\u0644\u0645\u0644\u0641 \u0645\u0641\u064a\u0647\u0648\u0634 \u062c\u062f\u0648\u0644"]

    table = doc.tables[0]
    rows = table.rows

    if len(rows) < 2:
        return [], ["\u0627\u0644\u062c\u062f\u0648\u0644 \u0644\u0627\u0632\u0645 \u064a\u0643\u0648\u0646 \u0641\u064a\u0647 \u0639\u0644\u0649 \u0627\u0644\u0623\u0642\u0644 \u0633\u0637\u0631 \u0639\u0646\u0627\u0648\u064a\u0646 \u0648\u0633\u0637\u0631 \u0628\u064a\u0627\u0646\u0627\u062a"]

    headers = [cell.text.strip() for cell in rows[0].cells]

    # Map headers
    topic_number_col = None
    title_col = None
    video_path_col = None
    thumbnail_path_col = None
    platform_field_cols = []  # [(col_index, platform_id, field_name)]

    for i, h in enumerate(headers):
        h_stripped = h.strip()
        h_lower = h_stripped.lower()

        if h_lower in ('\u0631\u0642\u0645', 'topic_number', '#'):
            topic_number_col = i
        elif h_lower in ('\u0639\u0646\u0648\u0627\u0646', 'title'):
            title_col = i
        elif h_lower in ('\u0641\u064a\u062f\u064a\u0648', 'video_path', '\u0645\u0633\u0627\u0631 \u0627\u0644\u0641\u064a\u062f\u064a\u0648'):
            video_path_col = i
        elif h_lower in ('\u0635\u0648\u0631\u0629', 'thumbnail', 'thumbnail_path', '\u0627\u0644\u0635\u0648\u0631\u0629 \u0627\u0644\u0645\u0635\u063a\u0631\u0629', '\u0645\u0633\u0627\u0631 \u0627\u0644\u0635\u0648\u0631\u0629'):
            thumbnail_path_col = i
        elif ' - ' in h_stripped or '.' in h_stripped:
            sep = ' - ' if ' - ' in h_stripped else '.'
            parts = h_stripped.split(sep, 1)
            if len(parts) == 2:
                plat_name = parts[0].strip()
                field_label = parts[1].strip()

                plat_id = _find_in_map(plat_name, platform_map)
                if plat_id is not None:
                    field_name = _find_in_map(field_label, field_map.get(plat_id, {}))
                    if field_name is None:
                        field_name = field_label
                    platform_field_cols.append((i, plat_id, field_name))

    if topic_number_col is None:
        return [], ["\u0645\u0641\u064a\u0634 \u0639\u0645\u0648\u062f '\u0631\u0642\u0645' \u0641\u064a \u0627\u0644\u062c\u062f\u0648\u0644"]

    MAX_ROWS = 1000
    if len(rows) - 1 > MAX_ROWS:
        return [], [f"\u0627\u0644\u062c\u062f\u0648\u0644 \u0641\u064a\u0647 {len(rows)-1} \u0633\u0637\u0631 \u2014 \u0627\u0644\u062d\u062f \u0627\u0644\u0623\u0642\u0635\u0649 {MAX_ROWS}"]

    # Parse data rows
    topics = []
    errors = []

    for row_idx, row in enumerate(rows[1:], start=2):
        cells = [cell.text.strip() for cell in row.cells]

        if not any(cells):
            continue

        # Topic number
        raw_num = cells[topic_number_col] if topic_number_col < len(cells) else ''
        try:
            topic_num = int(raw_num)
        except (ValueError, TypeError):
            errors.append(f"\u0633\u0637\u0631 {row_idx}: \u0631\u0642\u0645 \u0627\u0644\u0645\u0648\u0636\u0648\u0639 \u063a\u0644\u0637 '{raw_num}'")
            continue

        title = cells[title_col] if title_col is not None and title_col < len(cells) else None
        video_path = cells[video_path_col] if video_path_col is not None and video_path_col < len(cells) and cells[video_path_col] else None
        thumbnail_path = cells[thumbnail_path_col] if thumbnail_path_col is not None and thumbnail_path_col < len(cells) and cells[thumbnail_path_col] else None

        # Platform data
        platform_data_map = {}
        for col_idx, plat_id, field_name in platform_field_cols:
            if col_idx < len(cells) and cells[col_idx]:
                if plat_id not in platform_data_map:
                    platform_data_map[plat_id] = {}
                platform_data_map[plat_id][field_name] = cells[col_idx]

        topic = {
            "channel_id": channel_id,
            "topic_number": topic_num,
            "content_type": content_type,
            "title": title,
            "video_path": video_path,
            "thumbnail_path": thumbnail_path,
            "platform_data": [
                {"platform_id": pid, "field_values": fields}
                for pid, fields in platform_data_map.items()
            ] if platform_data_map else None,
        }
        topics.append(topic)

    return topics, errors


def generate_template(
    platforms_with_fields: List[Dict],
) -> bytes:
    """
    Generate an empty .docx template with correct headers.
    Returns bytes of the .docx file.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from io import BytesIO

    doc = Document()
    doc.core_properties.language = "ar"

    # Build headers
    headers = ['\u0631\u0642\u0645', '\u0639\u0646\u0648\u0627\u0646']
    for p in platforms_with_fields:
        plat_name = p['display_name'] or p['name']
        for f in p['fields']:
            label = f['field_label'] or f['field_name']
            headers.append(f"{plat_name} - {label}")

    # Create table
    table = doc.add_table(rows=2, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Example row
    table.rows[1].cells[0].text = '1'
    if len(headers) > 1:
        table.rows[1].cells[1].text = '\u0639\u0646\u0648\u0627\u0646 \u0627\u0644\u0641\u064a\u062f\u064a\u0648'

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _map_mg_field(fkey: str, fval: str, platform_map: Dict, field_map: Dict):
    """Map MG Ranner field key to (platform_id, field_name, value).
    MG Ranner keys: yt_title_1, yt_desc_1, yt_keywords, yt_screen, yt_thumb_1,
                     tt_desc, tt_screen, tt_title,
                     fb_desc, fb_screen, fb_title, fb_thumb, fb_keywords,
                     up_desc, up_screen,
                     tr_*  (translations — skip for now)
    """
    # Skip translation fields
    if fkey.startswith("tr_"):
        return None

    # Determine platform prefix → platform name
    prefix_map = {
        "yt": "youtube",
        "tt": "tiktok",
        "fb": "facebook",
        "up": "upscrolled",
    }
    prefix = fkey[:2]
    plat_name = prefix_map.get(prefix)
    if not plat_name:
        return None

    pid = _find_in_map(plat_name, platform_map)
    if pid is None:
        return None

    # Map field suffix to DB field name
    field_suffix = fkey[3:]  # e.g. "title_1", "desc_1", "keywords", "screen", "thumb_1"
    suffix_to_field = {
        "title_1": "title",
        "title_2": "title",  # fallback
        "title": "title",
        "desc_1": "description",
        "desc_2": "description",
        "desc": "description",
        "keywords": "tags",
        "screen": "screen_text",
        "thumb_1": "thumbnail_text",
        "thumb_2": "thumbnail_text",
        "thumb": "thumbnail_text",
    }

    fname = suffix_to_field.get(field_suffix)
    if not fname:
        return None

    # Verify field exists for this platform
    plat_fields = field_map.get(pid, {})
    actual_fname = _find_in_map(fname, plat_fields)
    if actual_fname is None:
        # Field doesn't exist for this platform — skip
        return None

    return (pid, actual_fname, fval)


def _find_in_map(key: str, mapping: Dict[str, any]) -> any:
    """Case-insensitive lookup in a dict."""
    if key in mapping:
        return mapping[key]
    key_lower = key.lower()
    for k, v in mapping.items():
        if k.lower() == key_lower:
            return v
    return None
